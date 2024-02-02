from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_protocol(doc_path, name, birth_number, sampling_date):
    # Vytvoření nového dokumentu
    doc = Document()

    # Přidání nadpisu
    title = doc.add_heading('Výsledný protokol genetického vyšetření', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Zarovnání na střed

    # Přidání informací do tabulky
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False  # Zakázat automatické nastavení šířky sloupců

    # Nastavení vlastní šířky sloupců
    for col in table.columns:
        col.width = Pt(200)

    # Přidání řádků a naplnění hodnotami
    rows = table.rows
    rows[0].cells[0].text = 'Jméno a příjmení:'
    rows[0].cells[1].text = name
    rows[1].cells[0].text = 'Rodné číslo:'
    rows[1].cells[1].text = birth_number
    rows[2].cells[0].text = 'Datum odběru:'
    rows[2].cells[1].text = sampling_date

    # Upravit formátování prvního sloupce (tučné písmo)
    for row in rows:
        if len(row.cells[0].paragraphs) > 0 and len(row.cells[0].paragraphs[0].runs) > 0:
            row.cells[0].paragraphs[0].runs[0].bold = True

    # Uložení dokumentu
    doc.save(doc_path)

# Příklad použití
create_protocol('Vysledny_Protokol.docx', 'Jan Novák', '850126/1158', '2024-02-01')